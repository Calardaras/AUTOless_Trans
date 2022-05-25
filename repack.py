import os,sys,re,docx
from docx import Document

path_org = "original"
path_otp = "output"
path_tmp = "process"
path_bfr = "tobetrans"
path_aft = "transed"

# 读取中间文件
key_f = open( path_tmp + '\\' + 'key.txt' ,'r',encoding="utf-8-sig") 
original_f = open( path_tmp + '\\' + 'value.txt' ,'r',encoding="utf-8-sig") 
key_l = list(key_f)
original_l = list(original_f)
key_f.close()
original_f.close()
contents_f = open( path_tmp + '\\' + 'contents.txt' ,'r',encoding="utf-8-sig") 


#处理目录和引索
file_name =[]
index = []
for content in contents_f:
    file_name.append(re.search('^[^,]+', content).group(0))
    index.append(re.search('[^,]+$', content).group(0))

contents_f.close()

formed_f = open( path_tmp + '\\' + 'form_value zh'+'.txt' ,'r',encoding="utf-8-sig") 
for root,dirs,files in os.walk(path_org):
    if re.search('(?<=english)[^.]+', root):
        root = 'output\\' + re.search('(?<=english)[^.]+', root).group(0)
        if not os.path.exists(root):
            os.makedirs(root)
i = 0 #文件引索
j = 0 #行数
curr_yml = open( path_otp  + '\\' + re.search('(?<=(\\\\english\\\\))[^.]+', file_name[i]).group(0) + '.yml','w+',encoding="utf-8-sig") 
print(path_otp  + '\\' + re.search('(?<=(\\\\english\\\\))[^.]+', file_name[i]).group(0) + '.yml')
for line in formed_f:
    if line != '':
        if j< int(index[i])-1:
            translation = re.sub('"','',line)
            translation = '"' + re.sub('\n','',translation) + '"' 
            curr_yml.write(re.search('^[^(\n)]+',key_l[j]).group(0) + ' ' + translation +'\n')
            j+=1
        else:
            curr_yml.close()
            i+=1
            curr_yml = open( path_otp  + '\\' + re.search('(?<=(\\\\english\\\\))[^.]+', file_name[i]).group(0) + '.yml','w+',encoding="utf-8-sig") 
            print(path_otp  + '\\' + re.search('(?<=(\\\\english\\\\))[^.]+', file_name[i]).group(0) + '.yml')
            translation = re.sub('"','',line)
            translation = '"' + re.sub('\n','',translation) + '"' 
            curr_yml.write(re.search('^[^(\n)]+',key_l[j]).group(0) + ' ' + translation +'\n')
            j+=1
# i = 0
# j = 0
# curr_yml = open( path_otp  + '\\' + re.search('(?<=(\\\\english\\\\))[^.]+', file_name[i]).group(0) + '.yml','w+',encoding="utf-8-sig") 
# for order in range(count_files(path_aft)):
#     curr_docx = Document(path_aft+'\\value_'+ str(order) +' zh.docx')
#     print('开始处理'+'value_'+ str(order) +' zh.docx')
#     for paragraph in curr_docx.paragraphs:
#         if formed_f != '':
#             if j< int(index[i])-1:
#                 translation = re.sub('"','',formed_f)
#                 translation = '"' + re.sub('\n','',translation) + '"' 
#                 curr_yml.write(re.search('^[^(\n)]+',key_l[j]).group(0) + ' ' + translation +'\n')
#                 j+=1
#             else:
#                 curr_yml.close()
#                 i+=1
#                 curr_yml = open( path_otp  + '\\' + re.search('(?<=(\\\\english\\\\))[^.]+', file_name[i]).group(0) + '.yml','w+',encoding="utf-8-sig") 
#                 translation = re.sub('"','',formed_f)
#                 translation = '"' + re.sub('\n','',translation) + '"' 
#                 curr_yml.write(re.search('^[^(\n)]+',key_l[j]).group(0) + ' ' + translation +'\n')
#                 j+=1
# curr_yml.close()

# for curr_docx_name in file_docx_name_list:
#     i = 0
#     j = 0
#     curr_docx = Document(curr_docx_name)
#     rename = re.sub('l_english','l_simp_chinese',file_name[i])
#     curr_yml = open( path_otp  + '\\' + re.search('(?<=(\\\\english\\\\))[^.]+', rename).group(0) + '.yml','w+',encoding="utf-8-sig") 
#     curr_yml.write('[\n')
#     for paragraph in curr_docx.paragraphs:
#         if formed_f != '':
#             if j< int(index[i])-1:
#                 translation = re.sub('"','',formed_f)
#                 translation = '"' + re.sub('\n','',translation) + '"' 
#                 original = re.sub('\n','',original_l[j])
#                 curr_yml.write(
#                     ' {\n'+
#                     '  \"key\": "' + re.search('^[^(\n)]+',key_l[j]).group(0) +'",\n'+
#                     '  \"original\":' + original +',\n'+
#                     '  \"translation\": ' + translation +',\n'+
#                     '  \"stage\":' + str(1) +'\n'+
#                     ' },\n'
#                 )
#                 j+=1
#             else:
#                 curr_yml.write(']')
#                 curr_yml.close()
#                 i+=1
#                 rename = re.sub('l_english','l_simp_chinese',file_name[i])
#                 curr_yml = open( path_otp  + '\\' + re.search('(?<=(\\\\english\\\\))[^.]+', rename).group(0) + '.yml','w+',encoding="utf-8-sig") 
#                 curr_yml.write('[\n')
#                 translation = re.sub('"','',formed_f)
#                 translation = '"' + re.sub('\n','',translation) + '"' 
#                 original = re.sub('\n','',original_l[j])
#                 curr_yml.write(
#                     ' {\n'+
#                     '  \"key\": "' + re.search('^[^(\n)]+',key_l[j]).group(0) +'",\n'+
#                     '  \"original\":' + original +',\n'+
#                     '  \"translation\": ' + translation +',\n'+
#                     '  \"stage\":' + str(1) +'\n'+
#                     ' },\n'
#                 )
#                 j+=1
#     curr_yml.write('{}]')
#     curr_yml.close()

                