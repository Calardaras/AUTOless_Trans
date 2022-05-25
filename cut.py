import os,sys,re,docx
from docx import Document

path_org = "original"
path_otp = "output"
path_tmp = "process"
path_bfr = "tobetrans"
path_aft = "transed"

key_f = open( path_tmp + '\\' + 'key.txt' ,'r',encoding="utf-8-sig") 
value_f = open( path_tmp + '\\' + 'norm_value.txt' ,'r',encoding="utf-8-sig") 
contents_f = open( path_tmp + '\\' + 'contents.txt' ,'r',encoding="utf-8-sig") 

i = 0
count = 0
curr_f = open( path_bfr + '\\' + 'value_'+ str(i) +'.txt' ,'w+',encoding="utf-8-sig") 
curr_d = docx.Document()
for line in value_f :
    re_line_re = re.search('[^"]+?(?=")',line)
    if re_line_re:
        re_line = re_line_re.group(0)
        re_line = re.sub( '□', '\n',re_line) +'\n========\n'
    else:
        re_line = '\n========\n'
    if count + len(re_line) < 950000:  #文档字符限制 默认九十五万
        curr_f.write(re_line)
        curr_d.add_paragraph(re_line)
        count += len(re_line)
    else:
        print(count)
        curr_d.save(path_bfr + '\\' + 'value_'+ str(i) +'.docx')
        curr_d = docx.Document()
        i += 1
        curr_f = open( path_bfr + '\\' + 'value_'+ str(i) +'.txt' ,'w+',encoding="utf-8-sig") 
        count = 0
        curr_f.write(re_line)
        curr_d.run(re_line)
        count += len(re_line)
curr_d.save(path_bfr + '\\' + 'value_'+ str(i) +'.docx')